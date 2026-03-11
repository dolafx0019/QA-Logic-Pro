import os
import pypdf
import docx
from fastapi import UploadFile, HTTPException
from io import BytesIO

MAX_FILE_SIZE_MB = 5
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024
MAX_TEXT_LENGTH = 150000  # Reasonable limit (~30k words) for MVP

class ParsedDocument:
    def __init__(self, text: str, file_name: str, file_type: str, warnings: list[str], parsing_notes: str):
        self.text = text
        self.file_name = file_name
        self.file_type = file_type
        self.warnings = warnings
        self.parsing_notes = parsing_notes

async def parse_uploaded_document(file: UploadFile) -> ParsedDocument:
    """
    Validates and extracts raw text from an uploaded document.
    Rejects unsupported files, oversized files, or scanned PDFs.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="Missing filename")
        
    ext = os.path.splitext(file.filename)[1].lower()
    allowed_extensions = {".pdf", ".docx", ".txt"}
    
    if ext not in allowed_extensions:
        raise HTTPException(status_code=400, detail=f"Unsupported file type '{ext}'. Allowed: PDF, DOCX, TXT")
        
    # Read file content
    content = await file.read()
    
    if len(content) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(status_code=413, detail=f"File exceeds maximum size of {MAX_FILE_SIZE_MB}MB")
        
    if len(content) == 0:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")
        
    warnings = []
    parsing_notes = ""
    text = ""
    
    try:
        if ext == ".txt":
            # Attempt to decode TXT
            try:
                text = content.decode('utf-8')
            except UnicodeDecodeError:
                text = content.decode('latin-1', errors='replace')
                warnings.append("Text file was not UTF-8, characters may be corrupted.")
            parsing_notes = "Extracted as plain text."
            
        elif ext == ".docx":
            # Parse DOCX
            try:
                doc = docx.Document(BytesIO(content))
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                parsing_notes = "Extracted from Word document."
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Failed to parse DOCX: {str(e)}")
                
        elif ext == ".pdf":
            # Parse PDF
            try:
                reader = pypdf.PdfReader(BytesIO(content))
                pages_text = []
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        pages_text.append(page_text)
                text = "\n".join(pages_text)
                
                num_pages = len(reader.pages)
                # Heuristic for scanned / image-only PDFs
                if num_pages > 0 and len(text.strip()) < num_pages * 50:
                    raise HTTPException(
                        status_code=422,
                        detail="PDF appears to be scanned or image-only. OCR is not supported in this phase. Please upload a text-based PDF."
                    )
                parsing_notes = f"Extracted from {num_pages}-page PDF."
                
            except HTTPException:
                raise
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Failed to parse PDF: {str(e)}")
                
        # Post-validation
        text = text.strip()
        if not text:
            raise HTTPException(status_code=422, detail="No readable text could be extracted from the document.")
            
        if len(text) > MAX_TEXT_LENGTH:
            text = text[:MAX_TEXT_LENGTH]
            warnings.append(f"Document text was truncated to the first {MAX_TEXT_LENGTH} characters to stay within AI context limits.")
            parsing_notes += " Document text was truncated."
            
        return ParsedDocument(
            text=text,
            file_name=file.filename,
            file_type=ext[1:].upper(),
            warnings=warnings,
            parsing_notes=parsing_notes
        )
        
    finally:
        await file.seek(0)
